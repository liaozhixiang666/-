from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QPalette
#import qtawesome
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt,Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor  # 设置字体颜色
from PyQt5.QtWidgets import   QFileDialog, QMessageBox
from xlrd import open_workbook
import matplotlib.pyplot as plt
import os
import time
from _pydecimal import Context,ROUND_HALF_UP
plt.rcParams['font.sans-serif']=['SimHei']
signal1 = True
ps1 = ps2 =ps3= ks1=ks2=ks3=zh1=zh2=zh3=ks4=True
#类和函数
class Chengji:
    def __init__(self,cj):
        self.zy = cj[0]   #专业
        self.ps = cj[1]   #平时成绩
        self.ks = cj[2]   #期末成绩  考试成绩
        self.zh = cj[3]   #综合成绩  最终成绩
        self.lb = cj[4]    #学员类别
#列表处理函数
#求平均值
def avg(a):
    s = 0
    for i in range(len(a)):
        s += a[i]
    return round(s/len(a),1)
#及格率
def jige_rate(a):
    jige = 0
    for i in range(len(a)):
        if a[i]>= 60:
            jige += 1
    return round(jige*100/len(a),1)
#得分分布
def fenbu1(s):
    a = b = c = d = e = 0
    final = s
    for i in range(len(final)):
        if final[i] < 60:
            a = a + 1
        elif 60 <= final[i] < 70:
            b = b + 1
        elif 70 <= final[i] < 80:
            c = c + 1
        elif 80 <= final[i] < 90:
            d = d + 1
        elif 90 <= final[i] <= 100:
                e = e + 1
    return [a,b,c,d,e]
def fenbu2(s):
    n=sum(s)
    for i in range(5):
        s[i]=round(s[i]/n, 3)
    return s
#画分布图
def draw_fenbu(f,t):
    name11 = ['[0,60)', '[60,70)', '[70,80)', '[80,90)', '[90,100]']
    score = f
    n = sum(f)
    plt.figure(figsize=(9,6), dpi=80)
    x = list(range(0, 5))
    y = score
    plt.bar(x, y, width=0.5,  color='silver',ec='k')
    for a, b in zip(x, y):
        plt.text(a,
                    b + 0.1,
                    '{:.1%}'.format(b / n),
                    ha='center',
                    va='bottom',
                    fontsize=20)

    plt.xticks(range(5), name11, size=20)
    plt.yticks([])
#plt.xlabel('得分区间',fontsize = 25)
    plt.ylabel('比例',fontsize = 25)
    plt.title(t+'分布图',fontsize = 30)
    plt.savefig(t+'分布图.png')#要加个名字
    plt.close()
def draw_zys(f,zy,p_name):#专业分布情况、专业名、图的名字
    x = ['[0,60)', '[60,70)', '[70,80)', '[80,90)', '[90,100]']
    fmt = ['ko-','kD-.','k2--','k.:','kX-.','kv--','ko-.','k^-','kD-','kv-','k*-','ks-','kX-','k.-','k2-']
    plt.figure(figsize=(9,6), dpi=80)
    for i in range(len(f)):
        plt.plot(x,f[i],fmt[i],label = zy[i],markersize=10)        
    plt.xticks(size=20)
    plt.yticks(size=20)
    plt.legend(fontsize = 13)
    plt.title(p_name,fontsize = 20)
    #plt.show()
    plt.savefig(p_name+'.png')
    plt.close()
def draw_zy(zy,avg,p_name):
    plt.figure(figsize=(9,6), dpi=80)
    plt.barh(zy,avg,color = 'silver')
    for a,b in zip(avg,zy):
        plt.text(a,b,'{:}'.format(a),ha='left',va='center',fontsize=20)
        plt.text(a,b,b,ha='right',va='center',fontsize=20)
    plt.xticks(size=20)
    plt.yticks([])
    plt.title(p_name,fontsize = 20)
    #plt.show()
    plt.savefig(p_name+'.png')
    plt.close()
#平均分表格
def pjf(zy,a):
    s = 0
    for i in range(len(a)):
        s += a[i]
    b = str(round(s/len(a),1))
    return [zy,str(len(a)),b,str(max(a)),str(min(a))]

#小题
class Xiaoti:
    def __init__(self,name):
        self.name = name
        self.a = []#得分率
        self.b = []#分值
        self.num = 0#数目
    def get_data(self,list1,fen):
        self.a.append(list1)
        self.b.append(fen)
        self.num += 1
    def draw(self):
        n = len(self.a)
        num = []
        for i in range(n):
            num.append('题'+str(i+1)) 
        x = list(range(n))
        y = self.a
        plt.figure(figsize=(n+1,6), dpi=80)
        plt.bar(x, y, color='silver',ec='k')
        for a, b in zip(x, y):
            plt.text(a,
                     b + 0.1,
                     str(b)+'%',
                     ha='center',
                     va='bottom',
                     fontsize=20)
        plt.xticks(x, num, size=20)
        #plt.xlabel('专业',fontsize = 25)
        plt.ylabel('平均得分',fontsize = 25)
        plt.title(self.name+'得分情况',fontsize = 30)
        plt.savefig(self.name+'得分情况.png')
        plt.close()
def defenlv(list1,fen):
    sum1 = 0
    n = len(list1)
    for i in range(n):
        sum1 += list1[i]
    return round(sum1 *100/ n/fen ,1)

class Ui_MainWindow2(object):
    def setupUi(self, MainWindow):
        self.file1 = ''
        if signal1:
            self.file2 = ''
        self.class_name = ''
        self.teacher_name = ''
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(582, 574)
        self.centralWidget = QtWidgets.QWidget(MainWindow)
        self.centralWidget.setObjectName("centralWidget")
        #
        self.layoutWidget = QtWidgets.QWidget(self.centralWidget)
        self.layoutWidget.setGeometry(QtCore.QRect(110, 71, 361, 41))
        self.layoutWidget.setObjectName("layoutWidget")
        self.horizontalLayout = QtWidgets.QHBoxLayout(self.layoutWidget)
        self.horizontalLayout.setContentsMargins(0, 0, 0, 0)
        self.horizontalLayout.setObjectName("horizontalLayout")
        #选择成绩分析文件
        self.label = QtWidgets.QLabel(self.layoutWidget)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(18)
        self.label.setFont(font)
        self.label.setObjectName("label")
        self.horizontalLayout.addWidget(self.label)
        #浏览1
        self.pushButton_browse = QtWidgets.QPushButton(self.layoutWidget)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(20)
        self.pushButton_browse.setFont(font)
        self.pushButton_browse.setObjectName("pushButton_browse")
        self.horizontalLayout.addWidget(self.pushButton_browse)
        #说明1
        self.pushButton_sm1 = QtWidgets.QPushButton(self.centralWidget)
        self.pushButton_sm1.setGeometry(QtCore.QRect(330, 120, 141, 41))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(18)
        self.pushButton_sm1.setFont(font)
        self.pushButton_sm1.setObjectName("pushButton_sm1")
        #
        self.horizontalLayoutWidget_2 = QtWidgets.QWidget(self.centralWidget)
        self.horizontalLayoutWidget_2.setGeometry(QtCore.QRect(110, 270, 361, 51))
        self.horizontalLayoutWidget_2.setObjectName("horizontalLayoutWidget_2")
        self.horizontalLayout_3 = QtWidgets.QHBoxLayout(self.horizontalLayoutWidget_2)
        self.horizontalLayout_3.setContentsMargins(0, 0, 0, 0)
        self.horizontalLayout_3.setObjectName("horizontalLayout_3")
        #课程名称
        self.class_name_2 = QtWidgets.QLabel(self.horizontalLayoutWidget_2)
        #self.class_name_2.setMinimumSize(QtCore.QSize(48, 49))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(18)
        self.class_name_2.setFont(font)
        self.class_name_2.setObjectName("class_name_2")
        self.horizontalLayout_3.addWidget(self.class_name_2)
        self.classname = QtWidgets.QLineEdit(self.horizontalLayoutWidget_2)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(18)
        self.classname.setFont(font)
        self.classname.setObjectName("classname")
        self.classname.textChanged[str].connect(self.onChanged2)
        #
        self.horizontalLayout_3.addWidget(self.classname)
        self.horizontalLayoutWidget_3 = QtWidgets.QWidget(self.centralWidget)
        self.horizontalLayoutWidget_3.setGeometry(QtCore.QRect(110, 330, 361, 41))
        self.horizontalLayoutWidget_3.setObjectName("horizontalLayoutWidget_3")
        self.horizontalLayout_4 = QtWidgets.QHBoxLayout(self.horizontalLayoutWidget_3)
        self.horizontalLayout_4.setContentsMargins(0, 0, 0, 0)
        self.horizontalLayout_4.setObjectName("horizontalLayout_4")
        #老师姓名
        self.label_2 = QtWidgets.QLabel(self.horizontalLayoutWidget_3)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(18)
        self.label_2.setFont(font)
        self.label_2.setObjectName("label_2")
        self.horizontalLayout_4.addWidget(self.label_2)
        self.teachername = QtWidgets.QLineEdit(self.horizontalLayoutWidget_3)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(18)
        self.teachername.setFont(font)
        self.teachername.setObjectName("teachername")
        self.teachername.textChanged[str].connect(self.onChanged3)
        #
        self.horizontalLayout_4.addWidget(self.teachername)
        self.horizontalLayoutWidget = QtWidgets.QWidget(self.centralWidget)
        self.horizontalLayoutWidget.setGeometry(QtCore.QRect(110, 381, 361, 151))
        self.horizontalLayoutWidget.setObjectName("horizontalLayoutWidget")
        self.horizontalLayout_2 = QtWidgets.QHBoxLayout(self.horizontalLayoutWidget)
        self.horizontalLayout_2.setContentsMargins(0, 0, 0, 0)
        self.horizontalLayout_2.setObjectName("horizontalLayout_2")
        #开始
        self.pushButton_start = QtWidgets.QPushButton(self.horizontalLayoutWidget)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(20)
        self.pushButton_start.setFont(font)
        self.pushButton_start.setObjectName("pushButton_start")
        self.horizontalLayout_2.addWidget(self.pushButton_start)
        self.class_nam = QtWidgets.QLabel(self.horizontalLayoutWidget)
        self.class_nam.setText("")
        self.class_nam.setObjectName("class_name")
        self.horizontalLayout_2.addWidget(self.class_nam)
        #返回
        self.pushButton_back = QtWidgets.QPushButton(self.horizontalLayoutWidget)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(20)
        self.pushButton_back.setFont(font)
        self.pushButton_back.setObjectName("pushButton_back")
        self.horizontalLayout_2.addWidget(self.pushButton_back)
        #
        if signal1:
            self.layoutWidget_2 = QtWidgets.QWidget(self.centralWidget)
            self.layoutWidget_2.setGeometry(QtCore.QRect(110, 170, 361, 41))
            self.layoutWidget_2.setObjectName("layoutWidget_2")
            self.horizontalLayout_6 = QtWidgets.QHBoxLayout(self.layoutWidget_2)
            self.horizontalLayout_6.setContentsMargins(0, 0, 0, 0)
            self.horizontalLayout_6.setObjectName("horizontalLayout_6")
            #选择题目分析文件
            self.label_4 = QtWidgets.QLabel(self.layoutWidget_2)
            font = QtGui.QFont()
            font.setFamily("Arial")
            font.setPointSize(18)
            self.label_4.setFont(font)
            self.label_4.setObjectName("label_4")
            self.horizontalLayout_6.addWidget(self.label_4)
            self.pushButton_browse_2 = QtWidgets.QPushButton(self.layoutWidget_2)
            font = QtGui.QFont()
            font.setFamily("Arial")
            font.setPointSize(20)
            self.pushButton_browse_2.setFont(font)
            self.pushButton_browse_2.setObjectName("pushButton_browse_2")
            self.horizontalLayout_6.addWidget(self.pushButton_browse_2)     
            #格式说明 按钮
            self.pushButton_sm2 = QtWidgets.QPushButton(self.centralWidget)
            self.pushButton_sm2.setGeometry(QtCore.QRect(330, 220, 141, 41))
            font = QtGui.QFont()
            font.setFamily("Arial")
            font.setPointSize(18)
            self.pushButton_sm2.setFont(font)
            self.pushButton_sm2.setObjectName("pushButton_sm2")
            
        MainWindow.setCentralWidget(self.centralWidget)

        self.retranslateUi(MainWindow)
        self.pushButton_browse.clicked.connect(lambda:self.browse1())
        if signal1:
            self.pushButton_browse_2.clicked.connect(lambda:self.browse2())
            self.pushButton_sm2.clicked.connect(lambda:self.sm2())
        self.pushButton_start.clicked.connect(lambda:self.jm2())
        self.pushButton_sm1.clicked.connect(lambda:self.sm1())
        self.pushButton_back.clicked.connect(lambda:back())
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle('成绩分析报告自动生成工具')
        self.label.setText(_translate("MainWindow", "选择成绩分析文件："))
        self.pushButton_browse.setText(_translate("MainWindow", "浏览"))
        self.pushButton_browse.setStyleSheet('''QPushButton{background:#F76677;border-radius:15px;}
QPushButton:hover{background:red;}''')
        self.class_name_2.setText(_translate("MainWindow", "课程名称"))
        self.label_2.setText(_translate("MainWindow", "老师姓名"))
        self.pushButton_start.setText(_translate("MainWindow", "开始分析"))
        self.pushButton_start.setStyleSheet('''QPushButton{background:#6DDF6D;border-radius:15px;}
QPushButton:hover{background:green;}''')
        self.pushButton_back.setText(_translate("MainWindow", "返回"))
        self.pushButton_back.setStyleSheet('''QPushButton{background:#F76677;border-radius:15px;}
QPushButton:hover{background:red;}''')
        self.pushButton_sm1.setText(_translate("MainWindow", "格式说明"))
        self.pushButton_sm1.setStyleSheet('''QPushButton{border-radius:15px;}
QPushButton:hover{background:blue;}''')
        if signal1:
            self.label_4.setText(_translate("MainWindow", "选择题目分析文件："))
            self.pushButton_browse_2.setText(_translate("MainWindow", "浏览"))
            self.pushButton_browse_2.setStyleSheet('''QPushButton{background:#F7D674;border-radius:15px;}
QPushButton:hover{background:yellow;}''')
            self.pushButton_sm2.setText(_translate("MainWindow", "格式说明"))
            self.pushButton_sm2.setStyleSheet('''QPushButton{border-radius:15px;}
QPushButton:hover{background:pink;}''')
    def browse1(self):        
        openfile_name, _ = QFileDialog.getOpenFileName(None,'选择文件','','Excel files(*.xlsx , *.xls)')
        if openfile_name :
            msg = '选择了文件: <b>{}</b>作为成绩分析文件'.format(openfile_name)
            QMessageBox.about(None,'提示', msg)
            self.file1 = openfile_name
    def browse2(self):        
        openfile_name, _ = QFileDialog.getOpenFileName(None,'选择文件','','Excel files(*.xlsx , *.xls)')
        if openfile_name :
            msg = '选择了文件: <b>{}</b>作为题目分析文件'.format(openfile_name)
            QMessageBox.about(None,'提示', msg)
            self.file2 = openfile_name
    def sm1(self):
        os.system('cmd /c 成绩分析文件说明.docx')
    def sm2(self):
        os.system('cmd /c 题目分析文件说明.docx')
    def onChanged2(self, text):
        self.class_name =text
    def onChanged3(self, text):
        self.teacher_name =text
    def test(self):
        if self.class_name and self.teacher_name:
            print( self.class_name, self.teacher_name )
    def jm2(self):
        if self.file1  and self.class_name and self.teacher_name:
            #开始
            index = dict()
            #需要的数据
            class_name = self.class_name#课程名称############################################################
            teacher_name = self.teacher_name#老师名字##############################################################
         
            wb = open_workbook(self.file1)#文件来源#########################################################
         

            if signal1 :
                if self.file2 == '':
                    QMessageBox.about(None,'提示', '请选择小题分析文件')
                    return 
                
                wb2 = open_workbook(self.file2)#小题####################################################################################
                

            sh1 = wb.sheet_by_index(0)
            #获取列名
            for col1 in range(0,sh1.ncols):
                if sh1.cell(0,col1).value =='专业' :
                    index['专业'] = col1 
                if  sh1.cell(0,col1).value =='平时成绩':
                    index['平时成绩'] = col1
                if  sh1.cell(0,col1).value =='期末成绩':
                    index['考试成绩'] = col1
                if  sh1.cell(0,col1).value == '成绩':
                    index['综合成绩'] = col1
                if sh1.cell(0,col1).value == '备注':
                    index['备注'] = col1
                if sh1.cell(0,col1).value == '类别':
                    index['学员类别'] = col1
            hang = [str(sh1.cell_value(0, i)) for i in range(0, sh1.ncols)]   
            if '专业' not in hang or '平时成绩' not in hang or '期末成绩' not in hang or '成绩' not in hang or '备注' not in hang or '类别' not in hang:
                QMessageBox.about(None,'提示', '请检查成绩分析文件的格式')
                return
            #获取成绩数据
            data = dict()
            zhuanye = []
            leibie = []
            for i in range(1,sh1.nrows):
                #备注为 缓考 和 旷考 的情况
                if sh1.cell(i,index['备注']).value == '缓考' or sh1.cell(i,index['备注']).value == '旷考':
                    continue
                else :
                    cj = []
                    cj.append(sh1.cell_value(i, index['专业']))
                    cj.append(sh1.cell(i,index['平时成绩']).value)
                    cj.append(sh1.cell(i,index['考试成绩']).value)
                    cj.append(sh1.cell(i,index['综合成绩']).value)
                    cj.append(sh1.cell(i,index['学员类别']).value)
                    data[i] = Chengji(cj)
                    if sh1.cell_value(i, index['专业']) in zhuanye:
                        pass
                    else :
                        zhuanye.append(sh1.cell(i, index['专业']).value)
                    if sh1.cell_value(i, index['学员类别']) in leibie:
                        pass
                    else :
                        leibie.append(sh1.cell(i, index['学员类别']).value)
            #三个成绩
            pingshi = []
            kaoshi = []
            zonghe = []
            tmp_zy = []
            zhh = []
            fzh = []
            wjj = []
            lhpy = []
            zhihui=feizhihui=lianhepeiyang=wujunji=0
            for value in data.values():
                pingshi.append(value.ps)
                kaoshi.append(value.ks)
                zonghe.append(value.zh)
                tmp_zy.append(value.zy)
                if value.lb=='指挥类':
                    zhihui=zhihui+1
                    zhh.append(value.zh)
                if value.lb=='非指挥类':
                    feizhihui=feizhihui+1
                    fzh.append(value.zh)
                if value.lb=='联合培养':
                    lianhepeiyang=lianhepeiyang+1
                    lhpy.append(value.zh)
                if value.lb=='无军籍':
                    wujunji=wujunji+1
                    wjj.append(value.zh)
            ksdx = '×××学院'+'××级'+':'
            
            for i in range(len(zhuanye)):
                ksdx = ksdx + zhuanye[i]+'('+str(tmp_zy.count(zhuanye[i]))+'人)、'
            nan = jnan = z = jy = y =''
            if avg(kaoshi)<65:
                nan = '√'
            elif avg(kaoshi)<70:
                jnan = '√'
            elif avg(kaoshi)<80:
                z = '√'
            elif avg(kaoshi)<85:
                jy = '√'
            elif avg(kaoshi)>85:
                y = '√'

            #考核小结
            ksdx = ':'
            for i in range(len(zhuanye)):
                ksdx = ksdx + zhuanye[i]+'('+str(tmp_zy.count(zhuanye[i]))+'人)、'
            kaoshiduixiang = ksdx[:-1]#考试对象
            num_ying =sh1.nrows-1#应该参加考试的人
            num_shi = len(data)#实际参加考试的人
            nandu = '难         度：' + '难（' +  nan +'  ）' + '较难（' + jnan + '  ）   中（' + z +' ）   较易（' + jy +'  ）   易（'+ y +'  ）\n'#考试难度
            #成绩分布
            cjfb1 = '最高：'+str(max(kaoshi))+'分；   最低：'+str(min(kaoshi))+'分；   平均：'+str(avg(kaoshi))+'分；   及格率：'+str(jige_rate(kaoshi))+'%。'
            cjfb2 = '优秀(90-100分)人数:'+str(fenbu1(kaoshi)[4])+' '+'比例:'+str(round(fenbu1(kaoshi)[4]*100/len(kaoshi),1))+'%'
            cjfb3 = '良好(80-89分)人数:'+str(fenbu1(kaoshi)[3])+' '+'比例:'+str(round(fenbu1(kaoshi)[3]*100/len(kaoshi),1))+'%'
            cjfb4 = '中等(70-79分)人数:'+str(fenbu1(kaoshi)[2])+' '+'比例:'+str(round(fenbu1(kaoshi)[2]*100/len(kaoshi),1))+'%'
            cjfb5 = '及格(60-69分)人数:'+str(fenbu1(kaoshi)[1])+' '+'比例:'+str(round(fenbu1(kaoshi)[1]*100/len(kaoshi),1))+'%'
            cjfb6 = '不及格(<60分)人数:'+str(fenbu1(kaoshi)[0])+' '+'比例:'+str(round(fenbu1(kaoshi)[0]*100/len(kaoshi),1))+'%'
            #小题
            #小题的数据处理
            if signal1 and self.file1  and self.file2 and self.class_name and self.teacher_name:
                ti_name = []
                ti = []
                kg = False
                sh1 = wb2.sheet_by_index(0)
                hang = [str(sh1.cell_value(0, i)) for i in range(0, sh1.ncols)]
                for col1 in range(0,sh1.ncols):
                    if ' 'in sh1.cell(0,col1).value :#并上界面选择
                        if len(sh1.cell(0,col1).value.split())==3:
                            if sh1.cell(0,col1).value.split()[0] not in ti_name:
                                ti_name.append(sh1.cell(0,col1).value.split()[0])
                                ti.append(Xiaoti(sh1.cell(0,col1).value.split()[0]))
                            for i in range(len(ti)):
                                if  ti[i].name == sh1.cell(0,col1).value.split()[0]:
                                    list2 = [float(sh1.cell_value(i, col1)) for i in range(1, sh1.nrows)]
                                    #index = int(sh1.cell(0,col1).value.split()[1])
                                    fen = float(sh1.cell(0,col1).value.split()[2])
                                    list1 = defenlv(list2,fen)
                                    ti[i].get_data(list1,fen)
                        else :
                                QMessageBox.about(None,'提示', '请检查题目分析文件的格式')
                                return
                        kg = True
                if kg == False:
                    QMessageBox.about(None,'提示', '请检查题目分析文件的格式')
                    return
            #生成 考核小结
            def shengcheng1():
                
                document = Document()
                document.styles['Normal'].font.size=Pt(10.5)
                document.styles['Normal'].font.name = u'Times New Roman'
                document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                #黑体
                style = document.styles.add_style('black' ,WD_STYLE_TYPE.PARAGRAPH)
                style.font.size=Pt(10.5)
                style.font.name=u'黑体'
                style._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
                #标题
                p = document.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run(u'国防科技大学本科课程考核小结')
                run.font.size = Pt(16)
                run.font.name=u'黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
                table = document.add_table(rows=1, cols=40,style='Table Grid')
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False
                def merge(list1,row):
                    for i in range(len(list1)-1):
                        table.cell(row,list1[i]).merge(table.cell(row,list1[i+1]-1))
                        table.cell(row,list1[i]).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                        table.cell(row,list1[i]).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                #第一行
                row = 0 
                table.rows[row].height=Cm(0.9) 
                table.cell(row,0).text = '课程名称'
                table.cell(row,0).paragraphs[0].style = 'black'
                table.cell(row,6).text = class_name
                table.cell(row,28).text = '课程编号'
                table.cell(row,28).paragraphs[0].style = 'black'
                merge([0,6,28,33,40],row)
                #第二行
                table.add_row()
                row = row+1
                table.rows[row].height=Cm(0.9) 
                table.cell(row,0).text = '开课单位'
                table.cell(row,0).paragraphs[0].style = 'black'
                tun = table.cell(row,6).paragraphs[0].add_run(u'××学院××系（所、重点实验室）××教研室（研究室、实验室）')
                tun.font.color.rgb = RGBColor(255, 0, 0)
                merge([0,6],row)
                table.cell(row,6).merge(table.cell(row,39))
                table.cell(row,6).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                #第三行
                table.add_row()
                row = row+1
                table.cell(row,0).text = '考核对象'
                table.cell(row,0).paragraphs[0].style = 'black'
                merge([0,6],row)
                table.cell(row,6).merge(table.cell(row,39))
                tun = table.cell(row,6).paragraphs[0].add_run(u'××学院××级')
                tun.font.color.rgb = RGBColor(255, 0, 0)
                table.cell(row,6).paragraphs[0].add_run(kaoshiduixiang)
                #table.cell(row,6).text = kaoshiduixiang
                #第四行
                table.add_row()
                row = row+1
                table.rows[row].height=Cm(0.9) 
                table.cell(row,0).text = '终结性考核\n日期'
                table.cell(row,0).paragraphs[0].style = 'black'
                tun = table.cell(row,6).paragraphs[0].add_run(u'××××年×月×日')
                tun.font.color.rgb = RGBColor(255, 0, 0)
                table.cell(row,12).text = '终结性考\n核时间'
                table.cell(row,17).text = '分钟'
                table.cell(row,12).paragraphs[0].style = 'black'
                table.cell(row,22).text = '应考人数'
                table.cell(row,22).paragraphs[0].style = 'black'
                table.cell(row,27).text = str(num_ying)
                table.cell(row,31).text = '实考人数'
                table.cell(row,31).paragraphs[0].style = 'black'
                table.cell(row,36).text = str(num_shi)
                merge([0,6,12,17,22,27,31,36,40],row)
                table.cell(row,17).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
                #第五行
                table.add_row()
                row = row+1
                table.rows[row].height=Cm(0.9)
                table.cell(row,0).text = '考教分离'
                table.cell(row,0).paragraphs[0].style = 'black'
                table.cell(row,6).text ='是（）   否（）'
                table.cell(row,17).text = '终结性考核命题形式'
                table.cell(row,17).paragraphs[0].style = 'black'
                table.cell(row,22).text = '试题库（ ） 试卷库（ ） 其它（ ）'
                merge([0,6,17,22,40],row)
                #第六行
                table.add_row()
                row = row+1
                table.rows[row].height=Cm(0.9) 
                table.cell(row,0).text = '终结性考核\n阅卷形式'
                table.cell(row,0).paragraphs[0].style = 'black'
                table.cell(row,6).text = '单人阅卷（ ）   多人流水作业阅卷（ ）   多人独立阅卷（ ）   其它（  ）'
                merge([0,6,40],row)
                #第七行
                table.add_row()
                row = row+1
                table.rows[row].height=Cm(1.1) 
                table.cell(row,0).text = '终结性考核\n试题类型及\n其分数比例'
                table.cell(row,0).paragraphs[0].style = 'black'
                merge([0,6],row)
                table.cell(row,6).merge(table.cell(row,39))
                tun = table.cell(row,6).paragraphs[0].add_run(u'1）单选题xx道，每道xx分，共xx分，占xx%，内容涉及xxxx，主要考察xxxx；\n......')
                tun.font.color.rgb = RGBColor(255, 0, 0)
                #第八、九行
                table.add_row()
                row = row+1
                table.rows[row].height=Cm(1.1) 
                table.cell(row,0).text = '终结性考核\n内容分析'
                table.cell(row,0).paragraphs[0].style = 'black'
                merge([0,6],row)
                table.cell(row,6).merge(table.cell(row,39))
                table.cell(row,6).text = '覆  盖  面 ：宽（  ）   中（  ）   窄（  ）\n'+nandu+'题         量：>100（  ） 70-100（  ） 50-70（  ） <50  ( )\n'+'有无 错误：有（  ）   无（  ）'
                table.add_row()
                row = row+1
                table.rows[row].height=Cm(0.6) 
                #table.cell(row,0).text = '分\t析'
                table.cell(row,0).paragraphs[0].style = 'black'
                merge([0,6,40],row)
                table.cell(row,6).text = '根据平均分判断难度：<65（难）,<70（较难）,70-80（中）,>80(较易),>85(易)'
                table.cell(row-1,0).merge(table.cell(row,5))
                #第十行
                table.add_row()
                row = row+1
                table.rows[row].height=Cm(1.1) 
                table.cell(row,0).text = '考核组织\n情    况'
                table.cell(row,0).paragraphs[0].style = 'black'
                merge([0,6,40],row)
                #第十一到十四行
                table.add_row()
                row = row+1
                table.rows[row].height=Cm(0.9)
                table.cell(row,0).text = '考核成绩\n分    布'
                table.cell(row,0).paragraphs[0].style = 'black'
                table.cell(row,6).text = cjfb1
                merge([0,6],row)
                table.cell(row,6).merge(table.cell(row, 39))
                table.cell(row, 6).vertical_alignment = WD_ALIGN_VERTICAL.CENTER#垂直居中
                table.add_row()
                row = row+1
                table.rows[row].height=Cm(0.9)
                merge([0,6,23,40],row)
                table.add_row()
                row = row+1
                table.rows[row].height=Cm(0.9)
                merge([0,6,23,40],row)
                table.add_row()
                row = row+1
                table.rows[row].height=Cm(0.9)
                merge([0,6,23,40],row)
                table.cell(row-3,0).merge(table.cell(row,5))    
                table.cell(row-2,6).text = cjfb2
                table.cell(row-2,23).text = cjfb3
                table.cell(row-1,6).text = cjfb4
                table.cell(row-1,23).text = cjfb5
                table.cell(row,6).text = cjfb6        
                #第15行插入图表
                table.add_row()
                row = row+1
                table.cell(row,0).text = '终结性考核\n各类题型学\n员答题情况\n分\t析'
                table.cell(row,0).paragraphs[0].style = 'black'
                merge([0,6],row)
                table.cell(row,6).merge(table.cell(row,39))
                #图片和表格从这里开始
                par = 0
                #table.cell(row,6).paragraphs[par].add_run('一.平时成绩').bold=True
                if signal1 and self.file2:
                    table.cell(row,6).add_paragraph('各类题型得分情况')
                    par += 1
                    for i in range(len(ti)):
                        ti[i].draw()
                        p = table.cell(row,6).add_paragraph().add_run()
                        par += 1
                        p.add_picture(ti[i].name+'得分情况.png',width =Cm(len(ti[i].a)) ,height =Cm(6) )
                        os.remove(ti[i].name+'得分情况.png')
                        p = table.cell(row,6).add_paragraph()
                        par += 1
                        p.add_run('\t该图给出了'+ti[i].name+'各题的平均得分情况（平均得分=实际得分/该题分值×100%）。')
                        t = p.add_run('\t成绩分析内容······')
                        t.font.color.rgb = RGBColor(255, 0, 0)
                #倒数第三行
                table.add_row()
                row = row+1
                table.rows[row].height=Cm(7)
                table.cell(row,0).text = '学员考核\n成绩及学习\n情况分析'
                table.cell(row,0).paragraphs[0].style = 'black'
                merge([0,6],row)
                table.cell(row,6).merge(table.cell(row,39))
                par=0
                #考试成绩
                table.cell(row,6).add_paragraph()
                #par += 1
                #table.cell(row,6).paragraphs[par].add_run('二.期末成绩').bold=True
         
                if  ks1:
                    #table.cell(row,6).add_paragraph('期末成绩分布图')
                    #par += 1
                    draw_fenbu(fenbu1(kaoshi),'期末成绩')#######
                    p = table.cell(row,6).add_paragraph().add_run()
                    par += 1
                    p.add_picture('期末成绩分布图.png',width =Cm(8) ,height =Cm(6) )
                    os.remove('期末成绩分布图.png')
                    p = table.cell(row,6).add_paragraph()
                    par += 1
                    r = fenbu1(kaoshi)
                    p.add_run('\t考试成绩的分布情况为：[0,60)分段有'+str(r[0]) 
                                                        +'人，[60,70)分段有'+str(r[1])
                                                        +'人，[70,80)分段有'+str(r[2])
                                                        +'人，[80,90)分段有'+str(r[3])
                                                        +'人，[90,100]分段有'+str(r[4])+'人。')
                    run = p.add_run('补充分析内容······\n')
                    run.font.color.rgb = RGBColor(255, 0, 0)
                if  ks2:
                    p = table.cell(row,6).add_paragraph('各专业考试成绩平均分对比')
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    par += 1
                    table2 = table.cell(row,6).add_table(len(zhuanye)+1,5)
                    table2.style = 'Table Grid'
                    first = ['专业','人数','平均分','最高分','最低分']
                    for i in range(0,5):
                        if i == 0:
                            table2.columns[i].width=Cm(6)
                        else:
                            table2.columns[i].width=Cm(1.5)
                        table2.cell(0,i).text = first[i]
                        table2.cell(0,i).paragraphs[0].style = 'black'
                    for i in range(len(zhuanye)):
                        z = []
                        for value in data.values():
                            if value.zy == zhuanye[i]:
                                z.append(value.ks)###########
                        for j in range(5):
                            table2.cell(i+1,j).text = pjf(zhuanye[i],z)[j]
                    for i in range(5):
                        for j in range(len(zhuanye)+1):
                            table2.cell(j,i).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                            table2.cell(j,i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    if len(zhuanye)<16 :
                        zy_avg = []
                        for i in range(len(zhuanye)):
                            z = []
                            for value in data.values():
                                if value.zy == zhuanye[i]:
                                    z.append(value.ks)###############################################
                            s = avg(z)
                            zy_avg.append(s)
                        draw_zy(zhuanye,zy_avg,'各专业期末成绩平均分对比')
                        p = table.cell(row,6).add_paragraph().add_run()
                        par += 1
                        p.add_picture('各专业期末成绩平均分对比.png',width =Cm(8) ,height =Cm(6) )
                        os.remove('各专业期末成绩平均分对比.png')
                        p = table.cell(row,6).add_paragraph()
                        par += 1
                        p.add_run('\t该图展示了各专业期末成绩的平均分情况。')
                        t = p.add_run('\t成绩分析内容······')
                        t.font.color.rgb = RGBColor(255, 0, 0)
                if  ks3:
                    #新的作图
                    zy_f = []
                    for i in range(len(zhuanye)):
                        z = []
                        for value in data.values():
                            if value.zy == zhuanye[i]:
                                z.append(value.ks)###############################################
                        s = fenbu2(fenbu1(z))
                        zy_f.append(s)
                    draw_zys(zy_f,zhuanye,'各专业期末成绩分布情况')
                    p = table.cell(row,6).add_paragraph().add_run()
                    par += 1
                    p.add_picture('各专业期末成绩分布情况.png',width =Cm(8) ,height =Cm(6) )
                    os.remove('各专业期末成绩分布情况.png')
                    p = table.cell(row,6).add_paragraph()
                    par += 1
                    p.add_run('\t该图展示了各专业期末成绩的分布情况。')
                    t = p.add_run('\t成绩分析内容······')
                    t.font.color.rgb = RGBColor(255, 0, 0)
                if  0:   
                    table.cell(row,6).add_paragraph('各专业成绩分布情况')
                    par += 1
                    for i in range(len(zhuanye)):
                        z = []
                        for value in data.values():
                            if value.zy == zhuanye[i]:
                                z.append(value.ks)###
                        draw_fenbu(fenbu1(z),zhuanye[i]+'期末成绩')#######
                        #table.cell(row,6).add_paragraph(zhuanye[i]+'期末成绩分布图')
                        #par += 1
                        p = table.cell(row,6).add_paragraph().add_run()
                        par += 1
                        p.add_picture(zhuanye[i]+'期末成绩分布图.png',width =Cm(8) ,height =Cm(6) )
                        os.remove(zhuanye[i]+'期末成绩分布图.png')
                if  ps1:
                    #table.cell(row,6).add_paragraph('平时成绩分布图')
                    #par += 1
                    draw_fenbu(fenbu1(pingshi),'平时成绩')#######
                    p = table.cell(row,6).add_paragraph().add_run()
                    par += 1
                    p.add_picture('平时成绩分布图.png',width =Cm(8) ,height =Cm(6) )
                    os.remove('平时成绩分布图.png')
                    p = table.cell(row,6).add_paragraph()
                    par += 1
                    r = fenbu1(pingshi)
                    p.add_run('\t平时成绩的分布情况为：[0,60)分段有'+str(r[0]) 
                                                        +'人，[60,70)分段有'+str(r[1])
                                                        +'人，[70,80)分段有'+str(r[2])
                                                        +'人，[80,90)分段有'+str(r[3])
                                                        +'人，[90,100]分段有'+str(r[4])+'人。')
                    run = p.add_run('补充分析内容······\n')
                    run.font.color.rgb = RGBColor(255, 0, 0)
                if  ps2:
                    p = table.cell(row,6).add_paragraph('各专业平时成绩平均分对比')
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    par += 1
                    table2 = table.cell(row,6).add_table(len(zhuanye)+1,5)
                    table2.style = 'Table Grid'
                    first = ['专业','人数','平均分','最高分','最低分']
                    for i in range(0,5):
                        if i == 0:
                            table2.columns[i].width=Cm(6)
                        else:
                            table2.columns[i].width=Cm(1.5)
                        table2.cell(0,i).text = first[i]
                        table2.cell(0,i).paragraphs[0].style = 'black'
                    for i in range(len(zhuanye)):
                        z = []
                        for value in data.values():
                            if value.zy == zhuanye[i]:
                                z.append(value.ps)###########
                        for j in range(5):
                            table2.cell(i+1,j).text = pjf(zhuanye[i],z)[j]
                    for i in range(5):
                        for j in range(len(zhuanye)+1):
                            table2.cell(j,i).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                            table2.cell(j,i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    if len(zhuanye)<16 :
                        zy_avg = []
                        for i in range(len(zhuanye)):
                            z = []
                            for value in data.values():
                                if value.zy == zhuanye[i]:
                                    z.append(value.ps)###############################################
                            s = avg(z)
                            zy_avg.append(s)
                        draw_zy(zhuanye,zy_avg,'各专业平时成绩平均分对比')
                        p = table.cell(row,6).add_paragraph().add_run()
                        par += 1
                        p.add_picture('各专业平时成绩平均分对比.png',width =Cm(8) ,height =Cm(6) )
                        os.remove('各专业平时成绩平均分对比.png')
                        p = table.cell(row,6).add_paragraph()
                        par += 1
                        p.add_run('\t该图展示了各专业平时成绩的平均分情况。')
                        t = p.add_run('\t成绩分析内容······')
                        t.font.color.rgb = RGBColor(255, 0, 0)
                if  ps3:
                    #新的作图
                    zy_f = []
                    for i in range(len(zhuanye)):
                        z = []
                        for value in data.values():
                            if value.zy == zhuanye[i]:
                                z.append(value.ps)###############################################
                        s = fenbu2(fenbu1(z))
                        zy_f.append(s)
                    draw_zys(zy_f,zhuanye,'各专业平时成绩分布情况')
                    p = table.cell(row,6).add_paragraph().add_run()
                    par += 1
                    p.add_picture('各专业平时成绩分布情况.png',width =Cm(8) ,height =Cm(6) )
                    os.remove('各专业平时成绩分布情况.png')
                    p = table.cell(row,6).add_paragraph()
                    par += 1
                    p.add_run('\t该图展示了各专业的平时成绩分布情况。')
                    t = p.add_run('\t成绩分析内容······')
                    t.font.color.rgb = RGBColor(255, 0, 0)
                if  0:   
                    table.cell(row,6).add_paragraph('各专业成绩分布情况')
                    par += 1
                    for i in range(len(zhuanye)):
                        z = []
                        for value in data.values():
                            if value.zy == zhuanye[i]:
                                z.append(value.ps)
                        draw_fenbu(fenbu1(z),zhuanye[i]+'平时成绩')#######
                        #table.cell(row,6).add_paragraph(zhuanye[i]+'平时成绩分布图')
                        #par += 1
                        p = table.cell(row,6).add_paragraph().add_run()
                        par += 1
                        p.add_picture(zhuanye[i]+'平时成绩分布图.png',width =Cm(8) ,height =Cm(6) )
                        os.remove(zhuanye[i]+'平时成绩分布图.png')
                        p = table.cell(row,6).add_paragraph()
                        par += 1
                        r = fenbu1(kaoshi)
                        p.add_run('\t考试成绩的分布情况为：[0,60)分段有'+str(r[0]) 
                                                        +'人，[60,70)分段有'+str(r[1])
                                                        +'人，[70,80)分段有'+str(r[2])
                                                        +'人，[80,90)分段有'+str(r[3])
                                                        +'人，[90,100]分段有'+str(r[4])+'人。')
                        run = p.add_run('补充分析内容······')
                
                #综合成绩
                table.cell(row,6).add_paragraph()
                #par += 1
                #table.cell(row,6).paragraphs[par].add_run('三.综合成绩').bold=True

                if  zh1:
                    #table.cell(row,6).add_paragraph('综合成绩分布图')
                    #par += 1
                    draw_fenbu(fenbu1(zonghe),'综合成绩')#######
                    p = table.cell(row,6).add_paragraph().add_run()
                    par += 1
                    p.add_picture('综合成绩分布图.png',width =Cm(8) ,height =Cm(6) )
                    os.remove('综合成绩分布图.png')
                    p = table.cell(row,6).add_paragraph()
                    par += 1
                    r = fenbu1(zonghe)
                    p.add_run('\t综合成绩的分布情况为：[0,60)分段有'+str(r[0]) 
                                                        +'人，[60,70)分段有'+str(r[1])
                                                        +'人，[70,80)分段有'+str(r[2])
                                                        +'人，[80,90)分段有'+str(r[3])
                                                        +'人，[90,100]分段有'+str(r[4])+'人。')
                    run = p.add_run('补充分析内容······\n')
                    run.font.color.rgb = RGBColor(255, 0, 0)
                if zh2:
                    p = table.cell(row,6).add_paragraph('各专业综合成绩平均分对比')
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    par += 1
                    table2 = table.cell(row,6).add_table(len(zhuanye)+1,5)
                    table2.style = 'Table Grid'
                    first = ['专业','人数','平均分','最高分','最低分']
                    for i in range(0,5):
                        if i == 0:
                            table2.columns[i].width=Cm(6)
                        else:
                            table2.columns[i].width=Cm(1.5)
                        table2.cell(0,i).text = first[i]
                        table2.cell(0,i).paragraphs[0].style = 'black'
                    for i in range(len(zhuanye)):
                        z = []
                        for value in data.values():
                            if value.zy == zhuanye[i]:
                                z.append(value.zh)###########
                        for j in range(5):
                            table2.cell(i+1,j).text = pjf(zhuanye[i],z)[j]
                    for i in range(5):
                        for j in range(len(zhuanye)+1):
                            table2.cell(j,i).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                            table2.cell(j,i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    if len(zhuanye)<16 :
                        zy_avg = []
                        for i in range(len(zhuanye)):
                            z = []
                            for value in data.values():
                                if value.zy == zhuanye[i]:
                                    z.append(value.zh)###############################################
                            s = avg(z)
                            zy_avg.append(s)
                        draw_zy(zhuanye,zy_avg,'各专业综合成绩平均分对比')
                        p = table.cell(row,6).add_paragraph().add_run()
                        par += 1
                        p.add_picture('各专业综合成绩平均分对比.png',width =Cm(8) ,height =Cm(6) )
                        os.remove('各专业综合成绩平均分对比.png')
                        p = table.cell(row,6).add_paragraph()
                        par += 1
                        p.add_run('\t该图展示了各专业综合成绩的平均分情况。')
                        t = p.add_run('\t成绩分析内容······')
                        t.font.color.rgb = RGBColor(255, 0, 0)
                if  zh3:
                    #新的作图
                    zy_f = []
                    for i in range(len(zhuanye)):
                        z = []
                        for value in data.values():
                            if value.zy == zhuanye[i]:
                                z.append(value.zh)###############################################
                        s = fenbu2(fenbu1(z))
                        zy_f.append(s)
                    draw_zys(zy_f,zhuanye,'各专业综合成绩分布情况')
                    p = table.cell(row,6).add_paragraph().add_run()
                    par += 1
                    p.add_picture('各专业综合成绩分布情况.png',width =Cm(8) ,height =Cm(6) )
                    os.remove('各专业综合成绩分布情况.png')
                    p = table.cell(row,6).add_paragraph()
                    par += 1
                    p.add_run('\t该图展示了各专业综合成绩的分布情况。')
                    t = p.add_run('\t成绩分析内容······')
                    t.font.color.rgb = RGBColor(255, 0, 0)
                if  0:   
                    table.cell(row,6).add_paragraph('各专业成绩分布情况')
                    par += 1
                    for i in range(len(zhuanye)):
                        z = []
                        for value in data.values():
                            if value.zy == zhuanye[i]:
                                z.append(value.zh)###
                        draw_fenbu(fenbu1(z),zhuanye[i]+'综合成绩')#######
                        #table.cell(row,6).add_paragraph(zhuanye[i]+'综合成绩分布图')
                        #par += 1
                        p = table.cell(row,6).add_paragraph().add_run()
                        par += 1
                        p.add_picture(zhuanye[i]+'综合成绩分布图.png',width =Cm(8) ,height =Cm(6) )
                        os.remove(zhuanye[i]+'综合成绩分布图.png')
                
                #倒数第二行
                table.add_row()
                row = row+1
                table.rows[row].height=Cm(7)
                table.cell(row,0).text = '存在的问题\n及改进意见'
                table.cell(row,0).paragraphs[0].style = 'black'
                table.cell(row,0).paragraphs[0].paragraph_format.keep_with_next = True
                merge([0,6,40],row)
                #倒数第一行
                table.add_row()
                row = row+1
                table.cell(row,0).merge(table.cell(row,18))
                table.cell(row,0).text = '任课教员（签名）：'
                table.cell(row,0).paragraphs[0].paragraph_format.keep_with_next = True
                table.cell(row,0).add_paragraph('\n\n\n\n\n\n\n\n\n年\t月\t日').alignment=WD_ALIGN_PARAGRAPH.RIGHT
                table.cell(row,19).merge(table.cell(row,39))
                table.cell(row,19).text ='系（所、室）主任（签名）：'
                table.cell(row,19).paragraphs[0].paragraph_format.keep_with_next = True
                table.cell(row,19).add_paragraph('\n\n\n\n\n\n\n\n\n年\t月\t日').alignment=WD_ALIGN_PARAGRAPH.RIGHT
                p=document.add_paragraph('注：每门次课程应填写一份考核小结。')
                p.paragraph_format.space_before = Pt(8)
                time1=time.strftime('%Y{y}%m{m}%d{d} %H{h}%M{f}%S{s}').format(y='年',m='月',d='日',h='时',f='分',s='秒')
                tmp='《'+self.class_name+'》'+'考核小结-'+self.teacher_name+'-'+time1
                document.save(tmp+'.docx')
             
            #生成 成绩统计表
            
            def cjtj(list1,zhuanye):
                n = len(list1)
                jige = youxiu =sum1= 0
                for i in range(n):
                    if list1[i]>= 60:
                        jige =jige + 1
                    if list1[i]>= 90:
                        youxiu = youxiu + 1
                    sum1 = sum1 + list1[i]
                avg1 = str(round(sum1/n,4))
                avg=str(Context(prec=3,rounding=ROUND_HALF_UP).create_decimal(avg1))
                rate1 = str(round(jige *100 / n,1))+'%'
                rate2 = str(round(youxiu *100 / n,1))+'%'
                return ['青年学员',zhuanye,teacher_name,str(n),str(jige),rate1,str(youxiu),rate2,avg]
            def shengcheng2(): 

                document = Document()
                document.styles['Normal'].font.size=Pt(10.5)
                document.styles['Normal'].font.name = u'Times New Roman'
                document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                p = document.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                #标题
                run = p.add_run(u'国防科技大学本科课程考核成绩统计表')
                run.font.size = Pt(16)
                run.font.name=u'黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')

                #课程名称
                p = document.add_paragraph()
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run(u'课程名称：')
                run = p.add_run(class_name)
                #表格主体
                table = document.add_table(rows=1, cols=9,style='Table Grid')
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
  
                #单元格内容居中
                def center(row):
                    for i in range(9):        
                        #水平居中
                        table.cell(row,i).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                        #垂直居中
                        table.cell(row,i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                #第1行
                row = 0
              
                first = ['学员类别','学院专业','任课教员姓名','参考\n人数','及格\n人数','及格\n率%','优秀\n人数','优秀\n率%','平均\n成绩']
                for i in range(9):   
                    table.cell(row,i).text = first[i]
                center(row)
                #每一类别
                if  zhihui>0:#指挥类
                    for j in range(len(zhuanye)):#专业和类别之间的关系？？
                        signalz1=False
                        for value in data.values():
                            if value.zy == zhuanye[j] and value.lb=='指挥类':
                                list_z = []
                                signalz1=True
                                break
                        if signalz1:
                            for value in data.values():
                                    if value.zy == zhuanye[j] and value.lb=='指挥类':
                                        list_z.append(value.zh)###     
                            table.add_row()#
                            row=row+1
                            table.rows[row].height=Cm(1.25) #行高
                            for p in [ 2, 3, 4, 5, 6, 7, 8]:   
                                table.cell(row,p).text = cjtj(list_z,zhuanye[j])[p]
                            run2=table.cell(row,1).paragraphs[0].add_run(u'××学院××级')
                            run2.font.color.rgb = RGBColor(255, 0, 0)
                            tmp_zhuanye = cjtj(list_z,zhuanye[j])[1]
                            table.cell(row,1).paragraphs[0].add_run('\n'+tmp_zhuanye)
                            center(row)
                    table.add_row()#小合计
                    row=row+1
                    table.rows[row].height=Cm(1.25) #行高
                    table.cell(row,0).paragraphs[0].add_run(u'指挥类')
                    table.cell(row,1).merge(table.cell(row,2))
                    table.cell(row,1).text=('小\t计')
                    for q in range(3,9):   
                        table.cell(row,q).text = cjtj(zhh,'')[q]
                    table.cell(1,0).merge(table.cell(row,0))
                    center(row)
                    #非指挥类
                if  feizhihui>0:#指挥类
                    tmp1=0
                    for j in range(len(zhuanye)):#专业和类别之间的关系？？
                        signalz2=False
                        for value in data.values():
                            if value.zy == zhuanye[j] and value.lb=='非指挥类':
                                list_z = []
                                signalz2=True
                                break
                        if signalz2:
                            for value in data.values():
                                    if value.zy == zhuanye[j] and value.lb=='非指挥类':
                                        list_z.append(value.zh)###     
                            table.add_row()#
                            row=row+1
                            table.rows[row].height=Cm(1.25) #行高
                            tmp1=tmp1+1
                            for p in [ 2, 3, 4, 5, 6, 7, 8]:   
                                table.cell(row,p).text = cjtj(list_z,zhuanye[j])[p]
                            run2=table.cell(row,1).paragraphs[0].add_run(u'××学院××级')
                            run2.font.color.rgb = RGBColor(255, 0, 0)
                            tmp_zhuanye = cjtj(list_z,zhuanye[j])[1]
                            table.cell(row,1).paragraphs[0].add_run('\n'+tmp_zhuanye)
                            center(row)
                    table.add_row()#小合计      
                    row=row+1
                    table.rows[row].height=Cm(1.25) #行高
                    table.cell(row,0).paragraphs[0].add_run(u'非指挥类')
                    table.cell(row,1).merge(table.cell(row,2))
                    table.cell(row,1).text=('小\t计')
                    for q in range(3,9):   
                        table.cell(row,q).text = cjtj(fzh,'')[q]
                    table.cell(row-tmp1,0).merge(table.cell(row,0))
                    center(row)
                    #无军籍
                if  wujunji>0:
                    tmp2=0
                    for j in range(len(zhuanye)):#专业和类别之间的关系？？
                        signalz3=False
                        for value in data.values():
                            if value.zy == zhuanye[j] and value.lb=='无军籍':
                                list_z = []
                                signalz3=True
                                break
                        if signalz3:
                            for value in data.values():
                                    if value.zy == zhuanye[j] and value.lb=='无军籍':
                                        list_z.append(value.zh)###     
                            table.add_row()#
                            row=row+1
                            table.rows[row].height=Cm(1.25) #行高
                            tmp2=tmp2+1
                            for p in [ 2, 3, 4, 5, 6, 7, 8]:   
                                table.cell(row,p).text = cjtj(list_z,zhuanye[j])[p]
                            run2=table.cell(row,1).paragraphs[0].add_run(u'××学院××级')
                            run2.font.color.rgb = RGBColor(255, 0, 0)
                            tmp_zhuanye = cjtj(list_z,zhuanye[j])[1]
                            table.cell(row,1).paragraphs[0].add_run('\n'+tmp_zhuanye)
                            center(row)
                    table.add_row()#小合计      
                    row=row+1
                    table.rows[row].height=Cm(1.25) #行高
                    table.cell(row,0).paragraphs[0].add_run(u'无军籍')
                    table.cell(row,1).merge(table.cell(row,2))
                    table.cell(row,1).text=('小\t计')
                    for q in range(3,9):   
                        table.cell(row,q).text = cjtj(wjj,'')[q]
                    table.cell(row-tmp2,0).merge(table.cell(row,0))
                    center(row)
                #联合培养
                if  lianhepeiyang>0:
                    tmp3=0
                    for j in range(len(zhuanye)):#专业和类别之间的关系？？
                        signalz4=False
                        for value in data.values():
                            if value.zy == zhuanye[j] and value.lb=='联合培养':
                                list_z = []
                                signalz4=True
                                break
                        if signalz4:
                            for value in data.values():
                                    if value.zy == zhuanye[j] and value.lb=='联合培养':
                                        list_z.append(value.zh)###     
                            table.add_row()#
                            row=row+1
                            table.rows[row].height=Cm(1.25) #行高
                            tmp3=tmp3+1
                            for p in [ 2, 3, 4, 5, 6, 7, 8]:   
                                table.cell(row,p).text = cjtj(list_z,zhuanye[j])[p]
                            run2=table.cell(row,1).paragraphs[0].add_run(u'××学院××级')
                            run2.font.color.rgb = RGBColor(255, 0, 0)
                            tmp_zhuanye = cjtj(list_z,zhuanye[j])[1]
                            table.cell(row,1).paragraphs[0].add_run('\n'+tmp_zhuanye)
                            center(row)
                    table.add_row()#小合计      
                    row=row+1
                    table.rows[row].height=Cm(1.25) #行高
                    table.cell(row,0).paragraphs[0].add_run(u'联合培养')
                    table.cell(row,1).merge(table.cell(row,2))
                    table.cell(row,1).text=('小\t计')
                    for q in range(3,9):   
                        table.cell(row,q).text = cjtj(lhpy,'')[q]
                    table.cell(row-tmp3,0).merge(table.cell(row,0))
                    center(row)
                #合计
                table.add_row()
                row = row+1
                table.rows[row].height=Cm(1.25)
                table.cell(row,0).merge(table.cell(row,2))
                table.cell(row,1).text=('合\t计')
                for i in range(3,9):   
                        table.cell(row,i).text = cjtj(zonghe,'')[i]
                center(row)
                #倒数第1行
                table.add_row()
                row = row+1
                table.cell(row,0).merge(table.cell(row,3))
                table.cell(row,4).merge(table.cell(row,8))
                table.cell(row,0).text = '任课教员（签名）：'
                table.cell(row,0).add_paragraph('\n\n\n\n\n\n年\t月\t日').alignment=WD_ALIGN_PARAGRAPH.RIGHT
                table.cell(row,5).text ='系（所、室）主任（签名）：'
                table.cell(row,5).add_paragraph('\n\n\n\n\n\n年\t月\t日').alignment=WD_ALIGN_PARAGRAPH.RIGHT
                p = document.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                run = p.add_run(u'注：学员类别分别为指挥类、非指挥类、无军籍、联合培养等。')
                time2=time.strftime('%Y{y}%m{m}%d{d} %H{h}%M{f}%S{s}').format(y='年',m='月',d='日',h='时',f='分',s='秒')
                tmp2='《'+self.class_name+'》'+'成绩统计表-'+self.teacher_name+'-'+time2
                document.save(tmp2+'.docx')


            #用来测试
            if 1:
                shengcheng1()
                shengcheng2()
                QMessageBox.about(None,'提示', '已经成功生成分析文件！\n请注意，红色部分需要检查或修改，空白部分需要手动填写')
        if self.file1 == '':
            QMessageBox.about(None,'提示', '请选择成绩分析文件')
        if self.teacher_name == '':
            QMessageBox.about(None,'提示', '请输入教师姓名')
        if self.class_name == '':
            QMessageBox.about(None,'提示', '请输入课程名称')


class Ui_MainWindow(object):
    def setupUi(self, MainWindow):
        
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(606, 593)
        self.centralWidget = QtWidgets.QWidget(MainWindow)
        self.centralWidget.setObjectName("centralWidget")
        self.label = QtWidgets.QLabel(self.centralWidget)
        self.label.setGeometry(QtCore.QRect(230, 30, 54, 12))
        self.label.setText("")
        self.label.setObjectName("label")
        self.label_2 = QtWidgets.QLabel(self.centralWidget)
        self.label_2.setGeometry(QtCore.QRect(200, 60, 211, 31))
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(24)
        self.label_2.setFont(font)
        self.label_2.setObjectName("label_2")
        self.label_3 = QtWidgets.QLabel(self.centralWidget)
        self.label_3.setGeometry(QtCore.QRect(10, 180, 54, 12))
        self.label_3.setText("")
        self.label_3.setObjectName("label_3")
        self.pushButton = QtWidgets.QPushButton(self.centralWidget)
        self.pushButton.setGeometry(QtCore.QRect(340, 400, 111, 41))
        self.pushButton.setObjectName("pushButton")
        self.pushButton_next = QtWidgets.QPushButton(self.centralWidget)
        self.pushButton_next.setGeometry(QtCore.QRect(140, 400, 111, 41))
        self.pushButton_next.setObjectName("pushButton_next")
        self.layoutWidget = QtWidgets.QWidget(self.centralWidget)
        self.layoutWidget.setGeometry(QtCore.QRect(28, 180, 171, 161))
        self.layoutWidget.setObjectName("layoutWidget")
        self.verticalLayout = QtWidgets.QVBoxLayout(self.layoutWidget)
        self.verticalLayout.setContentsMargins(0, 0, 0, 0)
        self.verticalLayout.setObjectName("verticalLayout")
        self.checkBox_ps1 = QtWidgets.QCheckBox(self.layoutWidget)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(12)
        self.checkBox_ps1.setFont(font)
        self.checkBox_ps1.setObjectName("checkBox_ps1")
        self.checkBox_ps1.setChecked(True)
        self.verticalLayout.addWidget(self.checkBox_ps1)
        self.checkBox_ps2 = QtWidgets.QCheckBox(self.layoutWidget)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(12)
        self.checkBox_ps2.setFont(font)
        self.checkBox_ps2.setObjectName("checkBox_ps2")
        self.checkBox_ps2.setChecked(True)
        self.verticalLayout.addWidget(self.checkBox_ps2)
        self.checkBox_ps3 = QtWidgets.QCheckBox(self.layoutWidget)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(12)
        self.checkBox_ps3.setFont(font)
        self.checkBox_ps3.setObjectName("checkBox_ps3")
        self.checkBox_ps3.setChecked(True)
        self.verticalLayout.addWidget(self.checkBox_ps3)
        self.checkBox_psall = QtWidgets.QCheckBox(self.layoutWidget)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(12)
        self.checkBox_psall.setFont(font)
        self.checkBox_psall.setObjectName("checkBox_psall")
        self.checkBox_psall.setChecked(True)
        self.verticalLayout.addWidget(self.checkBox_psall)
        self.layoutWidget1 = QtWidgets.QWidget(self.centralWidget)
        self.layoutWidget1.setGeometry(QtCore.QRect(210, 183, 169, 161))
        self.layoutWidget1.setObjectName("layoutWidget1")
        self.verticalLayout_2 = QtWidgets.QVBoxLayout(self.layoutWidget1)
        self.verticalLayout_2.setContentsMargins(0, 0, 0, 0)
        self.verticalLayout_2.setObjectName("verticalLayout_2")
        self.checkBox_ks1 = QtWidgets.QCheckBox(self.layoutWidget1)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(12)
        self.checkBox_ks1.setFont(font)
        self.checkBox_ks1.setObjectName("checkBox_ks1")
        self.checkBox_ks1.setChecked(True)
        self.verticalLayout_2.addWidget(self.checkBox_ks1)
        self.checkBox_ks2 = QtWidgets.QCheckBox(self.layoutWidget1)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(12)
        self.checkBox_ks2.setFont(font)
        self.checkBox_ks2.setObjectName("checkBox_ks2")
        self.checkBox_ks2.setChecked(True)
        self.verticalLayout_2.addWidget(self.checkBox_ks2)
        self.checkBox_ks3 = QtWidgets.QCheckBox(self.layoutWidget1)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(12)
        self.checkBox_ks3.setFont(font)
        self.checkBox_ks3.setObjectName("checkBox_ks3")
        self.checkBox_ks3.setChecked(True)
        self.verticalLayout_2.addWidget(self.checkBox_ks3)
        self.checkBox_ks4 = QtWidgets.QCheckBox(self.layoutWidget1)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(12)
        self.checkBox_ks4.setFont(font)
        self.checkBox_ks4.setObjectName("checkBox_ks4")
        self.checkBox_ks4.setChecked(True)
        self.verticalLayout_2.addWidget(self.checkBox_ks4)
        self.checkBox_ksall = QtWidgets.QCheckBox(self.layoutWidget1)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(12)
        self.checkBox_ksall.setFont(font)
        self.checkBox_ksall.setObjectName("checkBox_ksall")
        self.checkBox_ksall.setChecked(True)
        self.verticalLayout_2.addWidget(self.checkBox_ksall)
        self.layoutWidget2 = QtWidgets.QWidget(self.centralWidget)
        self.layoutWidget2.setGeometry(QtCore.QRect(400, 180, 169, 161))
        self.layoutWidget2.setObjectName("layoutWidget2")
        self.verticalLayout_3 = QtWidgets.QVBoxLayout(self.layoutWidget2)
        self.verticalLayout_3.setContentsMargins(0, 0, 0, 0)
        self.verticalLayout_3.setObjectName("verticalLayout_3")
        self.checkBox_zh1 = QtWidgets.QCheckBox(self.layoutWidget2)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(12)
        self.checkBox_zh1.setFont(font)
        self.checkBox_zh1.setObjectName("checkBox_zh1")
        self.checkBox_zh1.setChecked(True)
        self.verticalLayout_3.addWidget(self.checkBox_zh1)
        self.checkBox_zh2 = QtWidgets.QCheckBox(self.layoutWidget2)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(12)
        self.checkBox_zh2.setFont(font)
        self.checkBox_zh2.setObjectName("checkBox_zh2")
        self.checkBox_zh2.setChecked(True)
        self.verticalLayout_3.addWidget(self.checkBox_zh2)
        self.checkBox_zh3 = QtWidgets.QCheckBox(self.layoutWidget2)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(12)
        self.checkBox_zh3.setFont(font)
        self.checkBox_zh3.setObjectName("checkBox_zh3")
        self.checkBox_zh3.setChecked(True)
        self.verticalLayout_3.addWidget(self.checkBox_zh3)
        self.checkBox_zhall = QtWidgets.QCheckBox(self.layoutWidget2)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(12)
        self.checkBox_zhall.setFont(font)
        self.checkBox_zhall.setObjectName("checkBox_zhall")
        self.checkBox_zhall.setChecked(True)
        self.verticalLayout_3.addWidget(self.checkBox_zhall)
        self.horizontalLayoutWidget = QtWidgets.QWidget(self.centralWidget)
        self.horizontalLayoutWidget.setGeometry(QtCore.QRect(30, 130, 551, 41))
        self.horizontalLayoutWidget.setObjectName("horizontalLayoutWidget")
        self.horizontalLayout = QtWidgets.QHBoxLayout(self.horizontalLayoutWidget)
        self.horizontalLayout.setContentsMargins(0, 0, 0, 0)
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.label_7 = QtWidgets.QLabel(self.horizontalLayoutWidget)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.label_7.setFont(font)
        self.label_7.setObjectName("label_7")
        self.horizontalLayout.addWidget(self.label_7)
        self.label_5 = QtWidgets.QLabel(self.horizontalLayoutWidget)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.label_5.setFont(font)
        self.label_5.setObjectName("label_5")
        self.horizontalLayout.addWidget(self.label_5)
        self.label_8 = QtWidgets.QLabel(self.horizontalLayoutWidget)
        font = QtGui.QFont()
        font.setFamily("Arial")
        font.setPointSize(14)
        self.label_8.setFont(font)
        self.label_8.setObjectName("label_8")
        self.horizontalLayout.addWidget(self.label_8)
        MainWindow.setCentralWidget(self.centralWidget)
        self.checkBox_ps1.stateChanged.connect(lambda: self.change_ps())
        self.checkBox_ps2.stateChanged.connect(lambda: self.change_ps())
        self.checkBox_ps3.stateChanged.connect(lambda: self.change_ps())
        self.checkBox_ks1.stateChanged.connect(lambda: self.change_ks())
        self.checkBox_ks2.stateChanged.connect(lambda: self.change_ks())
        self.checkBox_ks3.stateChanged.connect(lambda: self.change_ks())
        self.checkBox_ks4.stateChanged.connect(lambda: self.change_ks())
        self.checkBox_zh1.stateChanged.connect(lambda: self.change_zh())
        self.checkBox_zh2.stateChanged.connect(lambda: self.change_zh())
        self.checkBox_zh3.stateChanged.connect(lambda: self.change_zh())
        self.checkBox_psall.stateChanged.connect(lambda: self.change_psall())
        self.checkBox_ksall.stateChanged.connect(lambda: self.change_ksall())
        self.checkBox_zhall.stateChanged.connect(lambda: self.change_zhall())
        self.retranslateUi(MainWindow)
        self.pushButton.clicked.connect(MainWindow.close)
        self.pushButton_next.clicked.connect(lambda:self.next())
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle('成绩分析报告自动生成工具')
        MainWindow.setWindowOpacity(0.98)
        #MainWindow.setWindowFlag(QtCore.Qt.FramelessWindowHint)
        MainWindow.setAutoFillBackground(True)
        pe = QPalette()
        pe.setColor(QPalette.Window,Qt.lightGray)
        MainWindow.setPalette(pe)
        self.label_2.setText(_translate("MainWindow", "选择分析内容"))
        #self.label_2.setStyleSheet('''QLabel{color:white;font-size:30px;font-family:Roman times;}''')
        self.pushButton.setText(_translate("MainWindow", "退出"))
        self.pushButton.setStyleSheet('''QPushButton{font-size:22px;background:#F76677;border-radius:15px;}
QPushButton:hover{background:red;}''')
        self.pushButton_next.setText(_translate("MainWindow", "下一步"))
        self.pushButton_next.setStyleSheet('''QPushButton{font-size:22px;background:#6DDF6D;border-radius:15px;}
QPushButton:hover{background:green;}''')
        self.checkBox_ps1.setText(_translate("MainWindow", "总体成绩分布"))
        self.checkBox_ps2.setText(_translate("MainWindow", "对比各专业平均分"))
        self.checkBox_ps3.setText(_translate("MainWindow", "分专业统计成绩分布"))
        self.checkBox_psall.setText(_translate("MainWindow", "以上全部"))
        self.checkBox_ks1.setText(_translate("MainWindow", "总体成绩分布"))
        self.checkBox_ks2.setText(_translate("MainWindow", "对比各专业平均分"))
        self.checkBox_ks3.setText(_translate("MainWindow", "分专业统计成绩分布"))
        self.checkBox_ks4.setText(_translate("MainWindow", "小题得分情况\n（需要题目分析文件）"))
        self.checkBox_ksall.setText(_translate("MainWindow", "以上全部"))
        self.checkBox_zh1.setText(_translate("MainWindow", "总体成绩分布"))
        self.checkBox_zh2.setText(_translate("MainWindow", "对比各专业平均分"))
        self.checkBox_zh3.setText(_translate("MainWindow", "分专业统计成绩分布"))
        self.checkBox_zhall.setText(_translate("MainWindow", "以上全部"))
        self.label_7.setText(_translate("MainWindow", "平时成绩："))
        #self.label_7.setStyleSheet('''QLabel{color:blue;font-size:25px;font-family:Roman times;}''')
        self.label_5.setText(_translate("MainWindow", "考试成绩："))
        #self.label_5.setStyleSheet('''QLabel{color:yellow;font-size:25px;font-family:Roman times;}''')
        self.label_8.setText(_translate("MainWindow", "  综合成绩："))
        #self.label_8.setStyleSheet('''QLabel{color:pink;font-size:25px;font-family:Roman times;}''')
        
    def change_ks(self):
        if self.checkBox_ks1.isChecked() and self.checkBox_ks2.isChecked() and self.checkBox_ks3.isChecked() and self.checkBox_ks4.isChecked():
            self.checkBox_ksall.setChecked(True)
        else:
            self.checkBox_ksall.setChecked(False)
        global ks1, ks2, ks3
        ks1 = self.checkBox_ks1.isChecked()
        ks2 = self.checkBox_ks2.isChecked()
        ks3 = self.checkBox_ks3.isChecked()
    def change_ps(self):
        if self.checkBox_ps1.isChecked() and self.checkBox_ps2.isChecked() and self.checkBox_ps3.isChecked():
            self.checkBox_psall.setChecked(True)
        else:
            self.checkBox_psall.setChecked(False)
        global ps1 , ps2 , ps3
        ps1 = self.checkBox_ps1.isChecked()
        ps2 = self.checkBox_ps2.isChecked()
        ps3 = self.checkBox_ps3.isChecked()
    def change_zh(self):
        if self.checkBox_zh1.isChecked() and self.checkBox_zh2.isChecked() and self.checkBox_zh3.isChecked():
            self.checkBox_zhall.setChecked(True)
        else:
            self.checkBox_zhall.setChecked(False)
        global zh1, zh2, zh3
        zh1 = self.checkBox_zh1.isChecked()
        zh2 = self.checkBox_zh2.isChecked()
        zh3 = self.checkBox_zh3.isChecked()
    def change_psall(self):
        if self.checkBox_psall.isChecked():
            self.checkBox_ps1.setChecked(True)
            self.checkBox_ps2.setChecked(True)
            self.checkBox_ps3.setChecked(True)

    def change_zhall(self):
        if self.checkBox_zhall.isChecked():
            self.checkBox_zh1.setChecked(True)
            self.checkBox_zh2.setChecked(True)
            self.checkBox_zh3.setChecked(True)

    def change_ksall(self):
        if self.checkBox_ksall.isChecked():
            self.checkBox_ks1.setChecked(True)
            self.checkBox_ks2.setChecked(True)
            self.checkBox_ks3.setChecked(True)
            self.checkBox_ks4.setChecked(True)

    def next(self):
        global signal1
        if  self.checkBox_ks4.isChecked():
            signal1 = True
        else :
            signal1 = False
        ui2 = Ui_MainWindow2()    
        ui2.setupUi(MainWindow)    
        MainWindow.show()
        
def back():
    global signal1
    signal1 = True
    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    
if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec_())
